"""
portfolio_hedge_engine.py  v2.0
================================
Applies the Nifty 50 drawdown hedge analysis to YOUR actual portfolio.
Now features dynamic 100-DMA calculation via live market data.

HOW IT WORKS:
  1. You define your portfolio as a list of holdings (asset type, sector, % allocation)
  2. Each holding gets a Nifty Beta — the sensitivity of that holding to Nifty moves
  3. Beta-weighted Nifty Exposure is computed for the whole portfolio
  4. All 8 hedge strategies are re-scaled to YOUR portfolio's actual exposure
  5. If you provide a total portfolio value (optional), all outputs are in ₹ as well
"""

import os, sys, argparse, textwrap
from datetime import date
import warnings
import numpy as np
import pandas as pd
import yfinance as yf
from openpyxl import Workbook
from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side, GradientFill)
from openpyxl.utils import get_column_letter

# Suppress yfinance timezone warnings for a cleaner CLI
warnings.filterwarnings('ignore', category=FutureWarning)

# ══════════════════════════════════════════════════════════════
# LIVE DATA MODULE
# ══════════════════════════════════════════════════════════════
def fetch_live_nifty_status():
    """Dynamically calculates the Nifty 100-DMA and checks the hedge trigger."""
    print("  [Live API] Fetching Nifty 50 data for dynamic 100-DMA calculation...")
    try:
        # Fetch 6 months of daily data to ensure we have 100 trading sessions
        ticker = yf.Ticker("^NSEI")
        hist = ticker.history(period="6mo")
        
        if hist.empty or len(hist) < 100:
            print("  Warning: Not enough data fetched for 100-DMA calculation.")
            return None

        # Calculate the 100-Day Simple Moving Average
        hist['100_DMA'] = hist['Close'].rolling(window=100).mean()
        
        latest_close = hist['Close'].iloc[-1]
        latest_100dma = hist['100_DMA'].iloc[-1]
        
        # Calculate the percentage gap between current price and 100-DMA
        gap_pct = ((latest_close - latest_100dma) / latest_100dma) * 100
        
        # Determine if the tactical trigger is met (Nifty is below 100-DMA by at least 2%)
        trigger_active = gap_pct <= -2.0 

        return {
            "close": round(latest_close, 2),
            "dma100": round(latest_100dma, 2),
            "gap_pct": round(gap_pct, 2),
            "trigger_active": trigger_active
        }
    except Exception as e:
        print(f"  Warning: Failed to fetch live Nifty data: {e}")
        return None

# ══════════════════════════════════════════════════════════════
# DEFAULT BETA MAP  (used if user doesn't specify per holding)
# ══════════════════════════════════════════════════════════════
BETA_DEFAULTS = {
    "nifty_etf"         : 1.00,
    "large_cap_stock"   : 1.05,
    "mid_cap_stock"     : 1.35,
    "small_cap_stock"   : 1.55,
    "multi_cap_fund"    : 1.10,
    "flexi_cap_fund"    : 1.05,
    "sectoral_it"       : 1.20,
    "sectoral_banking"  : 1.15,
    "sectoral_pharma"   : 0.80,
    "sectoral_fmcg"     : 0.70,
    "sectoral_auto"     : 1.25,
    "sectoral_infra"    : 1.30,
    "sectoral_realty"   : 1.45,
    "sectoral_metals"   : 1.40,
    "debt_liquid"       : 0.05,
    "debt_shortterm"    : 0.08,
    "debt_gilt"         : 0.12,
    "gold_etf_sgb"      : -0.20,
    "gold_fund"         : -0.18,
    "reit_invit"        : 0.55,
    "international_fund": 0.35,
    "us_tech_fund"      : 0.30,
    "cash_fd"           : 0.00,
    "pms_aif"           : 1.10,
}

# ══════════════════════════════════════════════════════════════
# HEDGE STRATEGY RESULTS FROM BACK-TEST (30yr, v3.1)
# ══════════════════════════════════════════════════════════════
STRATEGIES = {
    "S1_ShortFutures": {
        "name"        : "Short Nifty Futures (100%)",
        "short_name"  : "Short Futures",
        "win_rate"    : 1.000,
        "avg_pnl_pct" : 22.40,   # trueNet avg over 17 episodes (multi-roll cost deducted)
        "worst_pct"   : 9.97,
        "best_pct"    : 58.06,
        "avg_cost_pct": 0.20,
        "monthly_drag": 0.20,
        "liquidity"   : 9,
        "simplicity"  : 8,
        "score"       : 76.3,    # recomputed: wr×0.25 + min(avg,25)×1.50 + liq + simp×0.80 + std term
        "requires"    : "Futures account (F&O enabled)",
        "trigger"     : "VIX > 28 (Strategy C)",
        "note"        : "Full offset but unlimited upside sacrifice. Needs 25% margin buffer.",
        "beta_scale"  : True,
    },
    "S2_ATM_Put": {
        "name"        : "ATM Protective Put",
        "short_name"  : "ATM Put",
        "win_rate"    : 0.905,   # 90.5% — 2 episodes negative after multi-roll premium cost
        "avg_pnl_pct" : 12.30,
        "worst_pct"   : -1.01,  # can go negative in short shallow drawdowns
        "best_pct"    : 39.33,
        "avg_cost_pct": 1.20,
        "monthly_drag": 1.20,
        "liquidity"   : 7,
        "simplicity"  : 6,
        "score"       : 55.8,
        "requires"    : "F&O account",
        "trigger"     : "2-of-3 signals (Strategy B)",
        "note"        : "Strong protection for known catalysts. Cost scales with VIX; not always positive in shallow drawdowns.",
        "beta_scale"  : True,
    },
    "S3_OTM5_Put": {
        "name"        : "5% OTM Protective Put",
        "short_name"  : "5% OTM Put",
        "win_rate"    : 0.952,   # 95.2%
        "avg_pnl_pct" : 10.94,
        "worst_pct"   : -0.07,
        "best_pct"    : 40.56,
        "avg_cost_pct": 0.70,
        "monthly_drag": 0.70,
        "liquidity"   : 6,
        "simplicity"  : 6,
        "score"       : 54.0,
        "requires"    : "F&O account",
        "trigger"     : "2-of-3 signals (Strategy B backup)",
        "note"        : "Cheaper than ATM. First 5% loss unhedged. Good for deep crashes.",
        "beta_scale"  : True,
    },
    "S4_Gold15": {
        "name"        : "Gold Allocation (15%)",
        "short_name"  : "Gold 15%",
        "win_rate"    : 0.688,
        "avg_pnl_pct" : 0.66,
        "worst_pct"   : -3.03,
        "best_pct"    : 5.25,
        "avg_cost_pct": 0.05,
        "monthly_drag": 0.05,
        "liquidity"   : 8,
        "simplicity"  : 9,
        "score"       : 38.0,
        "requires"    : "Demat (Gold ETF / SGB)",
        "trigger"     : "Always-on (Strategy A)",
        "note"        : "Positive in 69% of drawdowns. Low cost. Global crisis hedge.",
        "beta_scale"  : False,
    },
    "S5_USDINR10": {
        "name"        : "USD/INR Long (10%)",
        "short_name"  : "USD/INR 10%",
        "win_rate"    : 0.625,   # 62.5% — corrected from old 93.8%
        "avg_pnl_pct" : 0.22,
        "worst_pct"   : -0.44,
        "best_pct"    : 1.78,
        "avg_cost_pct": 0.10,
        "monthly_drag": 0.10,
        "liquidity"   : 7,
        "simplicity"  : 7,
        "score"       : 33.5,
        "requires"    : "International fund / USD FD",
        "trigger"     : "INR stress (USD/INR > 88)",
        "note"        : "Useful for FX-driven drawdowns. Limited standalone protection.",
        "beta_scale"  : False,
    },
    "S6_Debt30": {
        "name"        : "Debt Rotation (30%)",
        "short_name"  : "Debt 30%",
        "win_rate"    : 1.000,
        "avg_pnl_pct" : 7.44,
        "worst_pct"   : 3.20,
        "best_pct"    : 18.62,
        "avg_cost_pct": 0.30,
        "monthly_drag": 0.10,
        "liquidity"   : 9,
        "simplicity"  : 9,
        "score"       : 56.5,
        "requires"    : "Liquid / overnight debt fund",
        "trigger"     : "Always-on (Strategy A core)",
        "note"        : "100% win rate. Reduces equity beta + earns 6.5% p.a. Simplest hedge.",
        "beta_scale"  : False,
    },
    "S7_Collar": {
        "name"        : "Collar (ATM Put + OTM Call)",
        "short_name"  : "Collar",
        "win_rate"    : 1.000,
        "avg_pnl_pct" : 15.96,
        "worst_pct"   : 5.23,
        "best_pct"    : 45.03,
        "avg_cost_pct": 0.60,
        "monthly_drag": 0.60,
        "liquidity"   : 5,
        "simplicity"  : 4,
        "score"       : 60.1,
        "requires"    : "F&O account (2 legs)",
        "trigger"     : "Pre-known events (elections, budgets)",
        "note"        : "Sells upside to pay for put. 100% win rate. Ideal for event-driven crashes.",
        "beta_scale"  : True,
    },
    "S8_Combined": {
        "name"        : "Combined (Put 50% + Gold 15% + Debt 20%)",
        "short_name"  : "Combined",
        "win_rate"    : 1.000,
        "avg_pnl_pct" : 12.12,
        "worst_pct"   : 5.28,
        "best_pct"    : 33.76,
        "avg_cost_pct": 0.80,
        "monthly_drag": 0.80,
        "liquidity"   : 6,
        "simplicity"  : 5,
        "score"       : 56.8,
        "requires"    : "F&O + Demat + Debt fund",
        "trigger"     : "2-of-3 signals active",
        "note"        : "Diversified hedge. No single point of failure. Recommended for HNIs.",
        "beta_scale"  : True,
    },
}

# Historical episodes for per-episode scenario analysis
EPISODES = [
    {"label":"2000 Dot-com",    "dd":-51.36, "dur":588, "rec":818},
    {"label":"2004 Election",   "dd":-29.94, "dur":124, "rec":199},
    {"label":"2008 GFC",        "dd":-59.86, "dur":293, "rec":739},
    {"label":"2010 Euro Debt",  "dd":-28.01, "dur":410, "rec":684},
    {"label":"2015 China",      "dd":-22.52, "dur":359, "rec":383},
    {"label":"2018 IL&FS",      "dd":-10.17, "dur":53,  "rec":123},
    {"label":"2020 COVID",      "dd":-38.44, "dur":69,  "rec":231},
    {"label":"2024 FII Exodus", "dd":-15.77, "dur":159, "rec":304},
    {"label":"2026 Curr. DD",   "dd":-15.18, "dur":87,  "rec":None},
]

# Actual hedge P&L (net) per episode per strategy
# Values verified against dashboard.html PNL_MATRIX (May 2026 data-verified pass)
# Corrections applied vs prior version:
#   2010 Euro Debt  Gold: 5.77→6.55, USDINR: 1.97→2.31
#   2015 China      Gold: 2.04→2.07
#   2018 IL&FS      Gold: 0.35→0.07
#   2026 Curr. DD   Gold: 1.58→1.73, USDINR: 0.54→0.48
EPISODE_PNL = {
    "2000 Dot-com"   : [51.36, 48.75, 45.50,  None,  None, 18.55, 49.95, 36.74],
    "2004 Election"  : [29.94, 28.16, 24.44, -1.47,  0.02,  9.64, 29.36, 19.04],
    "2008 GFC"       : [59.86, 57.33, 54.06,  1.07,  2.68, 19.52, 58.53, 42.76],
    "2010 Euro Debt" : [28.01, 26.04, 22.57,  6.55,  2.31, 10.59, 28.01, 25.85],
    "2015 China"     : [22.52, 20.99, 17.31,  2.07,  1.10,  8.67, 22.52, 18.32],
    "2018 IL&FS"     : [10.17,  8.36,  4.82,  0.07,  0.24,  3.33, 10.17,  6.75],
    "2020 COVID"     : [38.44, 37.09, 33.30,  1.30,  0.69, 11.90, 38.44, 27.77],
    "2024 FII Exodus": [15.77, 14.63, 10.70,  2.07,  0.44,  5.58, 15.77, 13.11],
    "2026 Curr. DD"  : [15.18, 14.33, 10.16,  1.73,  0.48,  5.02, 15.18, 12.09],
}

STRAT_KEYS = list(STRATEGIES.keys())

# ══════════════════════════════════════════════════════════════
# SAMPLE PORTFOLIO TEMPLATE
# ══════════════════════════════════════════════════════════════
SAMPLE_PORTFOLIO = [
    ("Nifty BeES ETF",           "nifty_etf",         "Index",         20.0,   None),
    ("HDFC Bank",                "large_cap_stock",   "Banking",        8.0,   1.10),
    ("Infosys",                  "large_cap_stock",   "IT",             7.0,   1.15),
    ("Reliance Industries",      "large_cap_stock",   "Conglomerate",   6.0,   1.05),
    ("ICICI Pru Mid Cap Fund",   "mid_cap_stock",     "Multi-sector",   5.0,   None),
    ("SBI Small Cap Fund",       "small_cap_stock",   "Multi-sector",   5.0,   None),
    ("Titan Company",            "large_cap_stock",   "Consumer",       4.0,   0.95),
    ("Adani Ports",              "large_cap_stock",   "Infra",          4.0,   1.20),
    ("Nippon IT ETF",            "sectoral_it",       "IT",             3.0,   None),
    ("Kotak Flexicap Fund",      "flexi_cap_fund",    "Multi-sector",   5.0,   None),
    ("HDFC Liquid Fund",         "debt_liquid",       "Debt",          10.0,   None),
    ("SBI Short Term Debt Fund", "debt_shortterm",    "Debt",           5.0,   None),
    ("Sovereign Gold Bond 2028", "gold_etf_sgb",      "Gold",           5.0,   None),
    ("Nippon Gold ETF",          "gold_etf_sgb",      "Gold",           3.0,   None),
    ("Embassy REIT",             "reit_invit",        "Real Estate",    3.0,   None),
    ("Mirae US Tech Fund",       "us_tech_fund",      "International",  4.0,   None),
    ("Cash / FD",                "cash_fd",           "Cash",           3.0,   None),
]

def generate_template_csv(path: str):
    rows = []
    for name, atype, sector, alloc, beta in SAMPLE_PORTFOLIO:
        rows.append({
            "Holding_Name"   : name,
            "Asset_Type"     : atype,
            "Sector"         : sector,
            "Allocation_Pct" : alloc,
            "Beta_Override"  : beta if beta is not None else "",
        })
    df = pd.DataFrame(rows)
    df.to_csv(path, index=False)
    print(f"  Template portfolio saved: {path}")

def load_portfolio(path: str) -> pd.DataFrame:
    df = pd.read_csv(path)
    df.columns = [c.strip() for c in df.columns]

    total = df["Allocation_Pct"].sum()
    if abs(total - 100) > 0.5:
        print(f"  Warning: Allocations sum to {total:.1f}% — normalising to 100%")
        df["Allocation_Pct"] = df["Allocation_Pct"] / total * 100

    def get_beta(row):
        ov = row.get("Beta_Override", "")
        if ov != "" and not pd.isna(ov):
            try:
                return float(ov)
            except:
                pass
        atype = str(row.get("Asset_Type", "")).strip().lower()
        return BETA_DEFAULTS.get(atype, 1.00)

    df["Beta"] = df.apply(get_beta, axis=1)
    df["Weight"] = df["Allocation_Pct"] / 100
    df["Weighted_Beta"] = df["Weight"] * df["Beta"]
    return df

def compute_portfolio_metrics(df: pd.DataFrame, portfolio_value: float = None):
    portfolio_beta   = df["Weighted_Beta"].sum()
    equity_weight    = df[df["Beta"] > 0.3]["Weight"].sum()
    hedge_weight     = df[df["Beta"] <= 0]["Weight"].sum()
    debt_weight      = df[(df["Beta"] > 0) & (df["Beta"] < 0.3)]["Weight"].sum()
    sector_breakdown = df.groupby("Sector")["Allocation_Pct"].sum().sort_values(ascending=False)
    top3_pct         = df.nlargest(3, "Allocation_Pct")["Allocation_Pct"].sum()

    metrics = {
        "portfolio_beta"   : round(portfolio_beta, 3),
        "equity_weight"    : round(equity_weight * 100, 1),
        "debt_weight"      : round(debt_weight * 100, 1),
        "hedge_weight"     : round(hedge_weight * 100, 1), 
        "top3_concentration": round(top3_pct, 1),
        "sector_breakdown" : sector_breakdown,
        "portfolio_value"  : portfolio_value,
        "num_holdings"     : len(df),
    }
    return metrics

def compute_hedge_sizing(portfolio_beta: float, portfolio_value: float = None):
    results = {}
    for key, strat in STRATEGIES.items():
        is_beta_scaled = strat.get("beta_scale", True) 

        if is_beta_scaled:
            effective_coverage_pct = strat["avg_pnl_pct"] * portfolio_beta
            worst_pct              = strat["worst_pct"]   * portfolio_beta
            best_pct               = strat["best_pct"]    * portfolio_beta
            monthly_drag_pct       = strat["monthly_drag"] * portfolio_beta
            drag_pa                = strat["monthly_drag"] * 12 * portfolio_beta
        else:
            effective_coverage_pct = strat["avg_pnl_pct"]
            worst_pct              = strat["worst_pct"]
            best_pct               = strat["best_pct"]
            monthly_drag_pct       = strat["monthly_drag"]
            drag_pa                = strat["monthly_drag"] * 12

        nifty_equiv_pct = portfolio_beta * 100 

        r = {
            "strategy_name"           : strat["name"],
            "short_name"              : strat["short_name"],
            "portfolio_beta"          : portfolio_beta,
            "is_beta_scaled"          : is_beta_scaled,
            "nifty_equiv_exposure_pct": round(nifty_equiv_pct, 1),
            "avg_protection_pct"      : round(effective_coverage_pct, 2),
            "worst_protection_pct"    : round(worst_pct, 2),
            "best_protection_pct"     : round(best_pct, 2),
            "monthly_drag_pct"        : round(monthly_drag_pct, 3),
            "annual_drag_pct"         : round(drag_pa, 2),
            "win_rate_pct"            : round(strat["win_rate"] * 100, 1),
            "composite_score"         : strat["score"],
            "requires"                : strat["requires"],
            "trigger"                 : strat["trigger"],
            "note"                    : strat["note"],
        }
        if portfolio_value:
            nifty_notional = portfolio_value * portfolio_beta
            r["nifty_equiv_exposure_inr"] = round(nifty_notional, 0)
            r["avg_protection_inr"]        = round(portfolio_value * effective_coverage_pct / 100, 0)
            r["monthly_drag_inr"]          = round(portfolio_value * monthly_drag_pct / 100, 0)
            r["annual_drag_inr"]           = round(portfolio_value * drag_pa / 100, 0)

        results[key] = r
    return results

def compute_scenario_analysis(portfolio_beta: float, portfolio_value: float = None):
    rows = []
    strat_list = list(STRATEGIES.keys())
    for ep in EPISODES:
        label = ep["label"]
        dd    = ep["dd"]
        port_loss_pct = max(dd * portfolio_beta, -100.0)
        row = {
            "Episode"           : label,
            "Nifty_DD_Pct"      : dd,
            "Portfolio_Loss_Pct": round(port_loss_pct, 2),
            "Duration_Days"     : ep["dur"],
            "Recovery_Days"     : ep["rec"] if ep["rec"] else "Ongoing",
        }
        ep_pnl = EPISODE_PNL.get(label, [None]*8)
        for i, sk in enumerate(strat_list):
            strat        = STRATEGIES[sk]
            is_beta_scaled = strat.get("beta_scale", True)
            raw_pnl      = ep_pnl[i]

            if raw_pnl is None:
                pnl_scaled   = None
                net_position = None
                offset_pct   = None
            else:
                pnl_scaled   = raw_pnl * portfolio_beta if is_beta_scaled else raw_pnl
                net_position = port_loss_pct + pnl_scaled
                offset_pct   = (pnl_scaled / abs(port_loss_pct) * 100) if port_loss_pct != 0 else 0

            sn = strat["short_name"]
            row[f"{sn}_HedgePnL_Pct"] = round(pnl_scaled, 2)   if pnl_scaled   is not None else None
            row[f"{sn}_NetPos_Pct"]   = round(net_position, 2)  if net_position is not None else None
            row[f"{sn}_Offset_Pct"]   = round(offset_pct, 1) if offset_pct is not None else None
            if portfolio_value:
                row[f"{sn}_HedgePnL_INR"] = round(portfolio_value * pnl_scaled / 100, 0) if pnl_scaled is not None else None
        if portfolio_value:
            row["Portfolio_Loss_INR"] = round(portfolio_value * port_loss_pct / 100, 0)
        rows.append(row)
    return pd.DataFrame(rows)

# ══════════════════════════════════════════════════════════════
# EXCEL WORKBOOK BUILDER
# ══════════════════════════════════════════════════════════════
def fmt_inr(val):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "—"
    val = float(val)
    if abs(val) >= 1e7:
        return f"₹{val/1e7:.2f}Cr"
    elif abs(val) >= 1e5:
        return f"₹{val/1e5:.2f}L"
    else:
        return f"₹{val:,.0f}"

def _style_header(cell, bg="1E3A5F", fg="FFFFFF", bold=True, size=10, wrap=False):
    cell.font = Font(name="Calibri", bold=bold, color=fg, size=size)
    cell.fill = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(horizontal="center", vertical="center",
                                wrap_text=wrap)

def _border(cell, style="thin"):
    s = Side(style=style, color="D0D0D0")
    cell.border = Border(left=s, right=s, top=s, bottom=s)

def _num_cell(cell, value, fmt="0.00", bold=False, color="000000"):
    cell.value = value if not isinstance(value, str) else value
    cell.number_format = fmt
    cell.font = Font(name="Calibri", size=10, bold=bold, color=color)
    cell.alignment = Alignment(horizontal="center", vertical="center")
    _border(cell)

def build_excel(df_portfolio, metrics, hedge_sizing, scenario_df, output_path):
    wb = Workbook()

    # ── Sheet 1: Portfolio Holdings ──────────────────────────
    ws1 = wb.active
    ws1.title = "1_Portfolio Holdings"
    ws1.freeze_panes = "A3"

    ws1.merge_cells("A1:H1")
    t = ws1["A1"]
    t.value = "PORTFOLIO HEDGE ENGINE — Holdings & Beta Analysis"
    t.font  = Font(name="Calibri", bold=True, size=14, color="1E3A5F")
    t.alignment = Alignment(horizontal="left", vertical="center")
    ws1.row_dimensions[1].height = 28

    headers = ["#", "Holding Name", "Asset Type", "Sector",
               "Allocation %", "Beta", "Weighted Beta", "Notes"]
    col_w   = [4, 30, 22, 18, 14, 9, 15, 35]
    for j, (h, w) in enumerate(zip(headers, col_w), 1):
        c = ws1.cell(row=2, column=j, value=h)
        _style_header(c)
        ws1.column_dimensions[get_column_letter(j)].width = w

    for i, row in df_portfolio.iterrows():
        r = i + 3
        beta = row["Beta"]
        if beta < 0:
            bg = "E8F5E9"
        elif beta < 0.1:
            bg = "F5F5F5"
        else:
            bg = "FFFFFF"

        data = [i+1, row["Holding_Name"], row["Asset_Type"], row.get("Sector",""),
                row["Allocation_Pct"], row["Beta"], row["Weighted_Beta"],
                BETA_DEFAULTS.get(row["Asset_Type"],"Custom β")]
        for j, val in enumerate(data, 1):
            c = ws1.cell(row=r, column=j, value=val)
            c.font = Font(name="Calibri", size=10)
            c.fill = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal="center" if j != 2 else "left",
                                    vertical="center")
            _border(c)
            if j == 5:
                c.number_format = "0.0\"%\""
            elif j in (6, 7):
                c.number_format = "0.00"

    sr = len(df_portfolio) + 5
    ws1.merge_cells(f"A{sr}:B{sr}")
    ws1[f"A{sr}"].value = "PORTFOLIO SUMMARY"
    ws1[f"A{sr}"].font  = Font(name="Calibri", bold=True, size=11, color="1E3A5F")

    summary_data = [
        ("Portfolio Beta (β)",       f"{metrics['portfolio_beta']:.3f}"),
        ("Equity Exposure",          f"{metrics['equity_weight']:.1f}%"),
        ("Debt / Liquid",            f"{metrics['debt_weight']:.1f}%"),
        ("Gold / Inverse Hedge",     f"{metrics['hedge_weight']:.1f}%"),
        ("Top-3 Concentration",      f"{metrics['top3_concentration']:.1f}%"),
        ("Number of Holdings",       str(metrics["num_holdings"])),
        ("Portfolio Value",          fmt_inr(metrics["portfolio_value"]) if metrics["portfolio_value"] else "Not provided"),
    ]
    for k, (lbl, val) in enumerate(summary_data):
        row_n = sr + 1 + k
        c1 = ws1.cell(row=row_n, column=1, value=lbl)
        c2 = ws1.cell(row=row_n, column=2, value=val)
        c1.font = Font(name="Calibri", bold=True, size=10)
        c2.font = Font(name="Calibri", size=10, color="1E3A5F")
        for c in (c1, c2):
            c.alignment = Alignment(vertical="center")
            _border(c)

    # ── Sheet 2: Hedge Sizing Recommendations ────────────────
    ws2 = wb.create_sheet("2_Hedge Recommendations")
    ws2.freeze_panes = "A3"

    ws2.merge_cells("A1:K1")
    t2 = ws2["A1"]
    t2.value = (f"HEDGE SIZING RECOMMENDATIONS  |  Portfolio β = {metrics['portfolio_beta']:.3f}  |  "
                f"Nifty-Equivalent Exposure: {metrics['portfolio_beta']*100:.0f}% of portfolio")
    t2.font = Font(name="Calibri", bold=True, size=12, color="1E3A5F")
    t2.alignment = Alignment(horizontal="left", vertical="center")
    ws2.row_dimensions[1].height = 24

    h2 = ["Rank", "Strategy", "Win Rate", "Avg Hedge Gain", "Worst Case",
          "Best Case", "Monthly Drag", "Annual Drag", "Score", "Trigger", "Requires"]
    cw2 = [6, 32, 10, 16, 14, 14, 13, 13, 9, 30, 30]
    for j, (h, w) in enumerate(zip(h2, cw2), 1):
        c = ws2.cell(row=2, column=j, value=h)
        _style_header(c)
        ws2.column_dimensions[get_column_letter(j)].width = w

    sorted_keys = sorted(STRATEGIES.keys(), key=lambda k: STRATEGIES[k]["score"], reverse=True)
    for rank, sk in enumerate(sorted_keys, 1):
        r = rank + 2
        hs = hedge_sizing[sk]
        st = STRATEGIES[sk]
        pv = metrics["portfolio_value"]

        def pct_or_inr(pct_key, inr_key):
            if pv and inr_key in hs:
                return f"{hs[pct_key]:.2f}%  ({fmt_inr(hs[inr_key])})"
            return f"{hs[pct_key]:.2f}%"

        row_data = [
            rank,
            st["name"],
            f"{hs['win_rate_pct']:.0f}%",
            pct_or_inr("avg_protection_pct", "avg_protection_inr"),
            f"{hs['worst_protection_pct']:.2f}%",
            f"{hs['best_protection_pct']:.2f}%",
            pct_or_inr("monthly_drag_pct", "monthly_drag_inr"),
            f"{hs['annual_drag_pct']:.2f}%",
            hs["composite_score"],
            st["trigger"],
            st["requires"],
        ]
        bg = "E8F4FD" if rank <= 3 else ("FFF8E1" if rank <= 5 else "FFFFFF")
        for j, val in enumerate(row_data, 1):
            c = ws2.cell(row=r, column=j, value=val)
            c.font = Font(name="Calibri", size=10, bold=(rank <= 3))
            c.fill = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal="center" if j != 2 else "left",
                                    vertical="center", wrap_text=(j >= 10))
            _border(c)
            ws2.row_dimensions[r].height = 24

    pr = len(STRATEGIES) + 5
    ws2.merge_cells(f"A{pr}:K{pr}")
    ws2[f"A{pr}"].value = "RECOMMENDED PLAYBOOK FOR THIS PORTFOLIO"
    ws2[f"A{pr}"].font  = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    ws2[f"A{pr}"].fill  = PatternFill("solid", fgColor="1E3A5F")
    ws2[f"A{pr}"].alignment = Alignment(vertical="center")
    ws2.row_dimensions[pr].height = 20

    pb = metrics["portfolio_beta"]
    playbook = [
        ("STRATEGY A — Always-On",
         f"Keep 20% Liquid Debt + 10% Gold. Your β={pb:.2f} means each 10% Nifty fall = ~{10*pb:.1f}% portfolio loss."),
        ("STRATEGY B — Tactical",
         f"Trigger: 2-of-3 signals (VIX>20, Nifty<100-DMA by -2%, FII outflow). Buy ATM Nifty Puts on {pb*100:.0f}% notional."),
        ("STRATEGY C — Crisis",
         f"VIX > 28. Short Nifty Futures covering {pb*100:.0f}% of portfolio notional. Needs F&O account + 25% margin."),
    ]
    for k, (title, desc) in enumerate(playbook):
        rn = pr + 1 + k
        ws2.merge_cells(f"C{rn}:K{rn}")
        c1 = ws2.cell(row=rn, column=1, value=title)
        ws2.merge_cells(f"A{rn}:B{rn}")
        c1 = ws2.cell(row=rn, column=1, value=title)
        c1.font = Font(name="Calibri", bold=True, size=10, color="1E3A5F")
        c1.alignment = Alignment(vertical="center")
        c2 = ws2.cell(row=rn, column=3, value=desc)
        c2.font = Font(name="Calibri", size=10)
        c2.alignment = Alignment(vertical="center", wrap_text=True)
        ws2.row_dimensions[rn].height = 30
        for j in range(1, 12):
            _border(ws2.cell(row=rn, column=j))

    # ── Sheet 3: Scenario Analysis ────────────────────────────
    ws3 = wb.create_sheet("3_Scenario Analysis")
    ws3.freeze_panes = "C3"

    ws3.merge_cells("A1:E1")
    t3 = ws3["A1"]
    t3.value = (f"SCENARIO ANALYSIS — How Each Hedge Would Have Protected Your Portfolio  "
                f"(β={metrics['portfolio_beta']:.3f})")
    t3.font = Font(name="Calibri", bold=True, size=12, color="1E3A5F")
    t3.alignment = Alignment(horizontal="left", vertical="center")
    ws3.row_dimensions[1].height = 24

    base_hdrs = ["Episode", "Nifty DD%", "Your Portfolio Loss%"]
    if metrics["portfolio_value"]:
        base_hdrs += ["Portfolio Loss (₹)"]
    for sk in STRAT_KEYS:
        sn = STRATEGIES[sk]["short_name"]
        base_hdrs += [f"{sn} Hedge Gain%", f"{sn} Net Pos%", f"{sn} Offset%"]

    for j, h in enumerate(base_hdrs, 1):
        c = ws3.cell(row=2, column=j, value=h)
        _style_header(c, wrap=True)
        ws3.column_dimensions[get_column_letter(j)].width = 15 if j <= 4 else 13
    ws3.row_dimensions[2].height = 32

    for i, ep_row in scenario_df.iterrows():
        r = i + 3
        port_loss = ep_row["Portfolio_Loss_Pct"]
        row_vals = [ep_row["Episode"], ep_row["Nifty_DD_Pct"], port_loss]
        if metrics["portfolio_value"]:
            row_vals.append(ep_row.get("Portfolio_Loss_INR", ""))
        for sk in STRAT_KEYS:
            sn = STRATEGIES[sk]["short_name"]
            row_vals += [
                ep_row.get(f"{sn}_HedgePnL_Pct", ""),
                ep_row.get(f"{sn}_NetPos_Pct", ""),
                ep_row.get(f"{sn}_Offset_Pct", ""),
            ]

        for j, val in enumerate(row_vals, 1):
            c = ws3.cell(row=r, column=j, value=val)
            c.font = Font(name="Calibri", size=10)
            c.alignment = Alignment(horizontal="center", vertical="center")
            _border(c)
            if isinstance(val, (int, float)):
                if j == 3 or (metrics["portfolio_value"] and j == 4):
                    c.font = Font(name="Calibri", size=10, color="C62828", bold=True)
                elif j > 4 and (j - (5 if metrics["portfolio_value"] else 4)) % 3 == 1:
                    if val > 0:
                        c.font = Font(name="Calibri", size=10, color="2E7D32")
                    elif val < 0:
                        c.font = Font(name="Calibri", size=10, color="C62828")

    # ── Sheet 4: Strategy Detail ──────────────────────────────
    ws4 = wb.create_sheet("4_Strategy Detail")

    ws4.merge_cells("A1:F1")
    t4 = ws4["A1"]
    t4.value = "STRATEGY DEEP-DIVE — How Each Hedge Works"
    t4.font = Font(name="Calibri", bold=True, size=12, color="1E3A5F")
    t4.alignment = Alignment(horizontal="left", vertical="center")
    ws4.row_dimensions[1].height = 24

    col_w4 = [5, 30, 60, 25, 25, 25]
    h4 = ["#", "Strategy", "How it Works / When to Use", "Requires", "Trigger", "Monthly Cost"]
    for j, (h, w) in enumerate(zip(h4, col_w4), 1):
        c = ws4.cell(row=2, column=j, value=h)
        _style_header(c)
        ws4.column_dimensions[get_column_letter(j)].width = w

    for rank, sk in enumerate(sorted_keys, 1):
        r = rank + 2
        st = STRATEGIES[sk]
        row_data = [rank, st["name"], st["note"], st["requires"], st["trigger"],
                    f"{st['monthly_drag']:.2f}% / month"]
        bg = "E8F4FD" if rank <= 3 else "FFFFFF"
        for j, val in enumerate(row_data, 1):
            c = ws4.cell(row=r, column=j, value=val)
            c.font = Font(name="Calibri", size=10, bold=(rank <= 3 and j == 2))
            c.fill = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal="left" if j in (2,3,4,5) else "center",
                                    vertical="center", wrap_text=True)
            _border(c)
        ws4.row_dimensions[r].height = 40

    # ── Sheet 5: Instructions ─────────────────────────────────
    #ws5 = wb.create_sheet("5_How to Use")
    #ws5.column_dimensions["A"].width = 100

    #lines = [
    #    ("PORTFOLIO HEDGE ENGINE — USER GUIDE", 14, True, "1E3A5F"),
    #    ("", 10, False, "000000"),
    #    ("HOW TO UPDATE THIS FILE WITH YOUR ACTUAL PORTFOLIO", 11, True, "1E3A5F"),
    #    ("1. Go to sheet '1_Portfolio Holdings'", 10, False, "000000"),
    #    ("2. Replace the sample holdings with your own investments.", 10, False, "000000"),
    #    ("   - Asset_Type must match one of the types in the Beta Defaults table below.", 10, False, "000000"),
    #    ("   - Beta_Override: leave blank to use the default, or enter a custom β value.", 10, False, "000000"),
    #    ("   - Allocation_Pct: your % weight in that holding.", 10, False, "000000"),
    #    ("3. Re-run portfolio_hedge_engine.py with --file your_portfolio.csv to refresh all outputs.", 10, False, "000000"),
    #    ("", 10, False, "000000"),
    #    ("WHAT IS PORTFOLIO BETA?", 11, True, "1E3A5F"),
    #    ("  β = 1.0 → Your portfolio moves exactly with the Nifty.", 10, False, "000000"),
    #    ("  β = 1.2 → A 10% Nifty fall causes ~12% portfolio loss.", 10, False, "000000"),
    #    ("", 10, False, "000000"),
    #    ("BETA DEFAULTS TABLE", 11, True, "1E3A5F"),
    #]
    #for row_n, (text, size, bold, color) in enumerate(lines, 1):
    #    c = ws5.cell(row=row_n, column=1, value=text)
    #    c.font = Font(name="Calibri", size=size, bold=bold, color=color)
    #    c.alignment = Alignment(vertical="center", wrap_text=True)
    #    ws5.row_dimensions[row_n].height = 18

    #start_r = len(lines) + 2
    #ws5.cell(row=start_r, column=1, value="Asset Type").font = Font(bold=True)
    #ws5.cell(row=start_r, column=2, value="Default Beta").font = Font(bold=True)
    #ws5.column_dimensions["B"].width = 15
    #for k, (atype, beta) in enumerate(BETA_DEFAULTS.items(), 1):
    #    c1 = ws5.cell(row=start_r + k, column=1, value=atype)
    #    c2 = ws5.cell(row=start_r + k, column=2, value=beta)
    #    c1.font = Font(name="Calibri", size=10)
    #    c2.font = Font(name="Calibri", size=10,
    #                   color="2E7D32" if beta < 0 else ("1565C0" if beta < 0.1 else "000000"))

    # ── Sheet 5: Dual-Path Comparison (Harvey Ch.3) ──────────────
    ws5 = wb.create_sheet("5_Vol Targeting Paths")
    ws5.column_dimensions["A"].width = 32
    ws5.column_dimensions["B"].width = 38
    ws5.column_dimensions["C"].width = 38

    # Title
    ws5.merge_cells("A1:C1")
    t5 = ws5["A1"]
    t5.value = "VOLATILITY TARGETING — TWO-PATH COMPARISON  (Harvey Ch.3)"
    t5.font  = Font(name="Calibri", bold=True, size=13, color="1E3A5F")
    t5.alignment = Alignment(horizontal="left", vertical="center")
    ws5.row_dimensions[1].height = 28

    # Use neutral seeds if no live data context available
    _vix_e  = 18.5
    _dd_e   = 0.0
    eq_wt_e = metrics.get("equity_weight", 60.0)
    _pe     = compute_equity_reduction_path(
        metrics["portfolio_beta"], _vix_e, _dd_e, eq_wt_e, metrics["portfolio_value"]
    )

    # Column headers
    for col, val in [(1, ""), (2, f"Path 1 — {_pe['p1_label']}"), (3, f"Path 2 — {_pe['p2_label']}")]:
        c = ws5.cell(row=2, column=col, value=val)
        _style_header(c, bg="1E3A5F" if col > 1 else "2E4A7A")
        ws5.row_dimensions[2].height = 22

    # Sub-descriptor row
    ws5.cell(row=3, column=2, value=_pe["p1_descriptor"]).font = Font(name="Calibri", italic=True, size=9, color="444444")
    ws5.cell(row=3, column=3, value=_pe["p2_descriptor"]).font = Font(name="Calibri", italic=True, size=9, color="444444")
    ws5.row_dimensions[3].height = 18

    pv5 = metrics["portfolio_value"]

    def _e5(row, label, v1, v2, bold=False):
        c0 = ws5.cell(row=row, column=1, value=label)
        c0.font  = Font(name="Calibri", bold=True, size=10)
        c0.fill  = PatternFill("solid", fgColor="F0F4FA")
        c0.alignment = Alignment(vertical="center")
        _border(c0)
        for col, val in [(2, v1), (3, v2)]:
            c = ws5.cell(row=row, column=col, value=val)
            c.font  = Font(name="Calibri", size=10, bold=bold)
            c.fill  = PatternFill("solid", fgColor="FFFFFF")
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            _border(c)
        ws5.row_dimensions[row].height = 20

    p1_put_str = f"{_pe['p1_hedge_pct']:.0f}% of notional" + (f"  ({fmt_inr(pv5 * _pe['p1_hedge_pct'] / 100):.0f})" if pv5 else "")
    p2_put_str = f"{_pe['p2_hedge_pct']:.0f}% of notional" + (f"  ({fmt_inr(pv5 * _pe['p2_hedge_pct'] / 100):.0f})" if pv5 else "")
    p2_eq_str  = f"-{_pe['p2_equity_trim']:.0f}%  →  {_pe['p2_new_equity']:.1f}%" + (f"  (~{fmt_inr(pv5 * _pe['p2_equity_trim'] / 100)})" if pv5 else "")

    rows5 = [
        (4,  "Current Regime",      _pe["regime"],                         _pe["regime"],                        True),
        (5,  "Equity Exposure",     f"{_pe['p1_equity_pct']:.1f}% (unchanged)", p2_eq_str,                      False),
        (6,  "Portfolio Beta",      f"{_pe['p1_beta']:.3f} (unchanged)",   f"{_pe['p2_effective_beta']:.3f} (reduced)", False),
        (7,  "Put / Hedge Size",    p1_put_str,                             p2_put_str,                           False),
        (8,  "Monthly Put Drag",    f"~{_pe['p1_monthly_drag']:.2f}%",     f"~{_pe['p2_monthly_drag']:.2f}%",   False),
        (9,  "Opp. Cost / month",   "None",                                 f"~{_pe['p2_opp_cost_monthly']:.2f}% (equity upside forgone)", False),
        (10, "Total Monthly Cost",  f"~{_pe['p1_monthly_drag']:.2f}%",     f"~{_pe['p2_total_drag']:.2f}%",     True),
        (11, "Action",              _pe["p1_action"],                       _pe["p2_action"],                     False),
        (12, "Best For",            _pe["p1_best_for"],                     _pe["p2_best_for"],                   False),
        (13, "Trade-off",           _pe["p1_tradeoff"],                     _pe["p2_tradeoff"],                   False),
    ]
    for r, label, v1, v2, bold in rows5:
        _e5(r, label, v1, v2, bold)
        ws5.row_dimensions[r].height = 28

    # Note row
    ws5.merge_cells("A15:C15")
    note_c = ws5["A15"]
    note_c.value = ("NOTE: Path 2 adapts Harvey's continuous vol-scalar to a regime-based step function "
                    "tied to VIX thresholds — avoiding daily churn, STT, and capital gains friction for "
                    "retail Indian investors. Choose the path that matches your tax situation and rebalancing comfort.")
    note_c.font      = Font(name="Calibri", italic=True, size=9, color="666666")
    note_c.alignment = Alignment(wrap_text=True, vertical="center")
    ws5.row_dimensions[15].height = 40

    wb.save(output_path)
    print(f"  Excel saved: {output_path}")

# ══════════════════════════════════════════════════════════════
# TEXT REPORT
# ══════════════════════════════════════════════════════════════
def generate_text_report(metrics, hedge_sizing, scenario_df, output_path, live_nifty=None):
    pv = metrics["portfolio_value"]
    pb = metrics["portfolio_beta"]

    lines = [
        "═" * 68,
        "  PORTFOLIO HEDGE ENGINE — ANALYSIS REPORT",
        f"  Generated: {date.today().strftime('%d %B %Y')}",
        "═" * 68,
        "",
        "PORTFOLIO SUMMARY",
        f"  Portfolio Beta (β):       {pb:.3f}",
        f"  Equity Exposure:          {metrics['equity_weight']:.1f}%",
        f"  Debt / Liquid:            {metrics['debt_weight']:.1f}%",
        f"  Gold / Inverse Hedge:     {metrics['hedge_weight']:.1f}%",
        f"  Top-3 Concentration:      {metrics['top3_concentration']:.1f}%",
        f"  Holdings:                 {metrics['num_holdings']}",
        f"  Portfolio Value:          {fmt_inr(pv) if pv else 'Not provided (% outputs only)'}",
        "",
        "WHAT YOUR BETA MEANS",
        f"  A 10% Nifty fall  →  ~{10*pb:.1f}% portfolio loss",
        f"  A 20% Nifty fall  →  ~{20*pb:.1f}% portfolio loss",
        f"  A 40% Nifty fall  →  ~{40*pb:.1f}% portfolio loss",
        f"  A 59% Nifty fall  →  ~{min(59*pb,100):.1f}% portfolio loss (2008 worst case)",
        "",
        "SECTOR BREAKDOWN",
    ]
    for sector, pct in metrics["sector_breakdown"].items():
        lines.append(f"  {sector:<25} {pct:>6.1f}%")

    lines += ["", "═" * 68, "HEDGE RECOMMENDATIONS (ranked by composite score)", "═" * 68, ""]

    sorted_keys = sorted(STRATEGIES.keys(), key=lambda k: STRATEGIES[k]["score"], reverse=True)
    for rank, sk in enumerate(sorted_keys, 1):
        hs = hedge_sizing[sk]
        st = STRATEGIES[sk]
        lines.append(f"  #{rank}  {st['name']}")
        lines.append(f"       Win Rate:       {hs['win_rate_pct']:.0f}%")
        lines.append(f"       Avg Protection: {hs['avg_protection_pct']:.2f}%"
                     + (f"  ({fmt_inr(hs.get('avg_protection_inr'))})" if pv else ""))
        lines.append(f"       Worst / Best:   {hs['worst_protection_pct']:.2f}% / {hs['best_protection_pct']:.2f}%")
        lines.append(f"       Annual Drag:    {hs['annual_drag_pct']:.2f}%"
                     + (f"  ({fmt_inr(hs.get('annual_drag_inr'))})" if pv else ""))
        lines.append(f"       Trigger:        {st['trigger']}")
        lines.append(f"       Requires:       {st['requires']}")
        lines.append(f"       Note:           {textwrap.fill(st['note'], 55, subsequent_indent=' '*23)}")
        lines.append("")

    lines += ["═" * 68, "SCENARIO ANALYSIS — Key Historical Episodes", "═" * 68, ""]
    for _, ep_row in scenario_df.iterrows():
        lines.append(f"  {ep_row['Episode']}")
        lines.append(f"    Nifty Drawdown:      {ep_row['Nifty_DD_Pct']:.1f}%")
        lines.append(f"    Your Portfolio Loss: {ep_row['Portfolio_Loss_Pct']:.1f}%"
                     + (f"  ({fmt_inr(ep_row.get('Portfolio_Loss_INR'))})" if pv else ""))
        lines.append("    Best hedge offset:")
        offsets = {}
        for sk in STRAT_KEYS:
            sn  = STRATEGIES[sk]["short_name"]
            val = ep_row.get(f"{sn}_Offset_Pct", None)
            if val is not None:
                offsets[sn] = float(val)
        if offsets:
            best_sn = max(offsets, key=offsets.get)
            lines.append(f"      → {best_sn}: {offsets[best_sn]:.0f}% of loss offset")
        else:
            lines.append("      → N/A (no data for this episode)")
        lines.append("")

    # --- INJECTING LIVE 100-DMA LOGIC HERE ---
    lines += ["═" * 68, "LIVE MARKET STATUS (Dynamic 100-DMA)", "═" * 68, ""]
    
    if live_nifty:
        lines.append(f"  Nifty 50 Current:   {live_nifty['close']:,.2f}")
        lines.append(f"  Nifty 100-DMA:      {live_nifty['dma100']:,.2f}")
        lines.append(f"  Gap to 100-DMA:     {live_nifty['gap_pct']:+.2f}%")
        
        if live_nifty['trigger_active']:
            lines.append("  ► TRIGGER STATUS:   [ACTIVE] Nifty is > 2% below 100-DMA.")
            lines.append("                      Deploy Strategy B hedges immediately.")
        else:
            lines.append("  ► TRIGGER STATUS:   [INACTIVE] Nifty is holding above threshold.")
            lines.append("                      Do not overpay for ATM puts today.")
    else:
        lines.append("  [Live data unavailable. Run script with internet connection.]")
    
    lines.append("")
    # ----------------------------------------

    lines += [
        "═" * 68,
        "RECOMMENDED PLAYBOOK  (Magnitude Hedging v3.2)",
        "═" * 68,
        "",
        "  HOW THE v3.2 ENGINE WORKS:",
        "  ─────────────────────────────────────────────────────────────────",
        "  Stage 1 — Early Warning (Sigmoid Onset)",
        "    Fires when volatility (RV/VIX) and FPI outflows are elevated.",
        "    Uses a sigmoid curve so the hedge builds gradually, not in jumps.",
        "    Equation: Score = 0.48·RV + 0.25·DMA-Gap + 0.15·FPI + 0.12·VIX",
        "",
        "  Stage 2 — Active Drawdown (Linear Interaction)",
        "    Fires once a confirmed peak-to-trough drawdown is underway.",
        "    Equation: Alloc = 12.5 + 0.52·VIX + 0.48·DD% + 5.10·[VIX×DD%÷100] - Penalty",
        "    Penalty reduces the allocation the longer the drawdown runs",
        "    (prevents overpaying premium in prolonged slow-bleed drawdowns).",
        "",
        "  Stage 3 — Escalation Overrides (Hard Glass-Smash Rules)",
        "    DD ≥ 20%          → Minimum 75% hedge enforced",
        "    DD ≥ 15% + VIX>28 → Full 100% hedge",
        "    DD ≥ 25%          → Full 100% hedge unconditional",
        "",
        "  De-escalation Gate (Memory State)",
        "    Once a high hedge is set, it is NOT reduced until:",
        "    (a) 50% of the drawdown is recovered, AND",
        "    (b) VIX and RV have stayed below their thresholds for 10 days.",
        "    This prevents false re-entry after a dead-cat bounce.",
        "  ─────────────────────────────────────────────────────────────────",
        "",
        "  STRATEGY A (Always-On — no trigger required):",
        "    Keep 20% Liquid Debt + 10% Gold at all times.",
        f"    This immediately reduces your effective β from {pb:.2f} to ~{pb*0.70:.2f}.",
        "    Annual drag: ~0.15–0.30%.",
        "",
        "  STRATEGY B (Tactical — v3.2 Stage 1 or 2 active):",
        "    Trigger: 2-of-3 signals (VIX>20, Nifty<100-DMA by -2%, FPI net outflow).",
        f"    Your Nifty-equivalent exposure: {pb*100:.0f}% of portfolio notional.",
        "    The v3.2 engine determines exact put allocation (typically 20–75%)",
        "    based on live VIX, drawdown depth, and FPI flow — see dashboard for",
        "    the current Beta-Adjusted Target shown in the Stage indicator.",
        "    Hold for 1 month, roll if signals still active. Exit when VIX < 15.",
        "",
        "  STRATEGY C (Crisis — Stage 3 active / VIX > 28):",
        f"    Short Nifty Futures covering {pb*100:.0f}% notional.",
        "    Combine with Strategy B collars to reduce premium cost.",
        "    Never exceed 75% notional in short futures.",
        "",
    ]

    # ── Dual-Path Comparison (Harvey Ch.3 Vol Targeting) ─────────────────────
    _vix_for_path = 18.5
    _dd_for_path  = 0.0
    if live_nifty:
        _vix_for_path = live_nifty.get('vix', 18.5)
        _dd_for_path  = live_nifty.get('gap_pct', 0.0)

    eq_wt  = metrics.get('equity_weight', 60.0)
    _path  = compute_equity_reduction_path(pb, _vix_for_path, _dd_for_path, eq_wt, pv)

    lines += [
        "═" * 68,
        "VOLATILITY TARGETING — TWO-PATH COMPARISON  (Harvey Ch.3)",
        "═" * 68,
        f"  Current Regime : {_path['regime']}",
        f"  VIX            : {_path['vix']:.1f}  |  DD from Peak: {_path['dd_pct']:.1f}%",
        "",
        f"  PATH 1 — {_path['p1_label']}",
        f"  {_path['p1_descriptor']}",
        "  ─────────────────────────────────────────────────",
        f"  Equity Exposure  : {_path['p1_equity_pct']:.1f}% (unchanged)",
        f"  Portfolio Beta   : {_path['p1_beta']:.3f} (unchanged)",
        f"  Put Allocation   : {_path['p1_hedge_pct']:.0f}% of notional"
        + (f"  ({fmt_inr(_path.get('p1_put_notional_inr'))})" if pv and 'p1_put_notional_inr' in _path else ""),
        f"  Monthly Drag     : ~{_path['p1_monthly_drag']:.2f}%"
        + (f"  ({fmt_inr(_path.get('p1_monthly_drag_inr'))})" if pv and 'p1_monthly_drag_inr' in _path else ""),
        f"  Action           : {_path['p1_action']}",
        f"  Best For         : {_path['p1_best_for']}",
        f"  Trade-off        : {_path['p1_tradeoff']}",
        "",
        f"  PATH 2 — {_path['p2_label']}",
        f"  {_path['p2_descriptor']}",
        "  ─────────────────────────────────────────────────",
        f"  Equity Trim      : -{_path['p2_equity_trim']:.0f}% from {eq_wt:.1f}%  →  {_path['p2_new_equity']:.1f}%"
        + (f"  (sell approx. {fmt_inr(_path.get('p2_equity_sell_inr'))})" if pv and 'p2_equity_sell_inr' in _path else ""),
        f"  Effective Beta   : {_path['p2_effective_beta']:.3f} (reduced from {pb:.3f})",
        f"  Residual Put     : {_path['p2_hedge_pct']:.0f}% of notional"
        + (f"  ({fmt_inr(_path.get('p2_put_notional_inr'))})" if pv and 'p2_put_notional_inr' in _path else ""),
        f"  Put Drag/month   : ~{_path['p2_monthly_drag']:.2f}%",
        f"  Opp. Cost/month  : ~{_path['p2_opp_cost_monthly']:.2f}%  (forgone equity upside)",
        f"  Total Drag/month : ~{_path['p2_total_drag']:.2f}%",
        f"  Action           : {_path['p2_action']}",
        f"  Best For         : {_path['p2_best_for']}",
        f"  Trade-off        : {_path['p2_tradeoff']}",
        "",
        "  NOTE: Path 2 is adapted from Harvey's continuous vol-scalar to a regime-based",
        "  step function tied to VIX thresholds — avoiding daily rebalancing churn,",
        "  STT, and capital gains friction for retail Indian investors.",
        "",
        "═" * 68,
        "CAVEATS",
        "═" * 68,
        "  • All P&L figures are pre-tax (STT, CGT, GST excluded).",
        "  • Beta estimates are approximate; actual correlations vary by market regime.",
        "  • Annual drag figures assume the hedge is held 100% of the time.",
        "    In practice, tactical hedges (B/C) are deployed only when signals fire.",
        "  • Past back-test performance does not guarantee future results.",
        "  • Short Nifty Futures require a 25% cash margin buffer for MTM margin calls.",
        "  • Pre-2008 VIX is a 20-day realised volatility proxy, not the actual NSE VIX.",
        "═" * 68,
    ]

    if hasattr(output_path, 'write'):
        output_path.write("\n".join(lines))
    else:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
    print(f"  Report saved: {output_path}")
# ==============================================================
# word document builder
# ==============================================================
def generate_docx_report(metrics, hedge_sizing, scenario_df, output_path, live_nifty=None):
    from docx import Document
    from datetime import date

    doc = Document()
    
    # Helper for formatting INR locally
    def local_fmt_inr(val):
        if not val: return "—"
        if abs(val) >= 1e7: return f"₹{val/1e7:.2f}Cr"
        if abs(val) >= 1e5: return f"₹{val/1e5:.2f}L"
        return f"₹{val:,.0f}"

    pv = metrics.get("portfolio_value")
    pb = metrics.get("portfolio_beta", 1.0)

    # 1. Header
    doc.add_heading('INSTITUTIONAL RISK ENGINE — ANALYSIS REPORT', 0)
    doc.add_paragraph(f"Generated: {date.today().strftime('%d %B %Y')}")
    
    # 2. Summary
    doc.add_heading('PORTFOLIO SUMMARY', level=1)
    p1 = doc.add_paragraph()
    p1.add_run(f"Portfolio Beta (β): {pb:.3f}\n").bold = True
    p1.add_run(f"Equity Exposure: {metrics.get('equity_weight', 0):.1f}%\n")
    p1.add_run(f"Debt / Liquid: {metrics.get('debt_weight', 0):.1f}%\n")
    p1.add_run(f"Gold / Inverse: {metrics.get('hedge_weight', 0):.1f}%\n")
    p1.add_run(f"Holdings: {metrics.get('num_holdings', 0)}\n")
    if pv:
        p1.add_run(f"Portfolio Value: {local_fmt_inr(pv)}\n")
        
    # 3. Live Data
    doc.add_heading('LIVE MARKET STATUS (Dynamic 100-DMA)', level=1)
    p_live = doc.add_paragraph()
    if live_nifty:
        p_live.add_run(f"Nifty 50 Current: {live_nifty['close']:,.2f}\n")
        p_live.add_run(f"Nifty 100-DMA: {live_nifty['dma100']:,.2f}\n")
        p_live.add_run(f"Gap to 100-DMA: {live_nifty['gap_pct']:+.2f}%\n")
        if live_nifty['trigger_active']:
            p_live.add_run("► TRIGGER STATUS: [ACTIVE] Nifty is > 2% below 100-DMA. Deploy Strategy B.\n").bold = True
        else:
            p_live.add_run("► TRIGGER STATUS: [INACTIVE] Nifty is holding above threshold.\n")
    else:
        p_live.add_run("[Live data unavailable]")

    # 4. Hedge Recommendations
    doc.add_heading('HEDGE RECOMMENDATIONS', level=1)
    sorted_keys = sorted(STRATEGIES.keys(), key=lambda k: STRATEGIES[k]["score"], reverse=True)
    for rank, sk in enumerate(sorted_keys, 1):
        hs = hedge_sizing[sk]
        st = STRATEGIES[sk]
        p2 = doc.add_paragraph()
        p2.add_run(f"#{rank} {st['name']}\n").bold = True
        p2.add_run(f"Win Rate: {hs['win_rate_pct']:.0f}%\n")
        
        pnl_str = f"{hs['avg_protection_pct']:.2f}%"
        if pv and 'avg_protection_inr' in hs: 
            pnl_str += f" ({local_fmt_inr(hs['avg_protection_inr'])})"
        p2.add_run(f"Avg Protection: {pnl_str}\n")
        
        drag_str = f"{hs['annual_drag_pct']:.2f}%"
        if pv and 'annual_drag_inr' in hs: 
            drag_str += f" ({local_fmt_inr(hs['annual_drag_inr'])})"
        p2.add_run(f"Annual Drag: {drag_str}\n")
        p2.add_run(f"Trigger: {st['trigger']}\n")
        p2.add_run(f"Note: {st['note']}")

    # 5. v3.2 Engine Guide
    doc.add_heading('MAGNITUDE HEDGING v3.2 — STAGE GUIDE', level=1)
    p_v3 = doc.add_paragraph()
    p_v3.add_run('How the Engine Stages Work').bold = True
    p_v3.add_run(
        'Stage 1 — Early Warning (Sigmoid Onset): Fires when realised volatility and FPI outflows '
        'are elevated before a confirmed drawdown. The hedge builds gradually using a sigmoid curve. '
        'Score = 0.48·RV + 0.25·DMA-Gap + 0.15·FPI + 0.12·VIX.'
        'Stage 2 — Active Drawdown (Linear Interaction): Fires once a peak-to-trough drawdown is '
        'confirmed. Alloc = 12.5 + 0.52·VIX + 0.48·DD% + 5.10·[VIX×DD%÷100] - TimePenalty. '
        'The time penalty reduces the allocation in prolonged slow-bleed drawdowns.'
        'Stage 3 — Escalation Overrides: Hard glass-smash rules. DD≥20% forces minimum 75% hedge. '
        'DD≥15% with VIX>28 forces 100%. DD≥25% forces 100% unconditionally.'
        'De-escalation Gate: Once a high hedge is set, it is NOT reduced until (a) 50% of the '
        'drawdown is recovered AND (b) VIX and RV have stayed below thresholds for 10 consecutive days. '
        'This prevents false re-entry after a dead-cat bounce.'
    )

    # 6. Recommended Playbook
    doc.add_heading('RECOMMENDED PLAYBOOK', level=1)
    p_pb = doc.add_paragraph()
    p_pb.add_run('Strategy A — Always-On Core Defence').bold = True
    p_pb.add_run(
        f'Keep ≥20% Liquid Debt + 10% Gold at all times. '
        f'This reduces your effective β from {pb:.2f} to ~{pb*0.70:.2f}. Annual drag: ~0.15–0.30%.'
    )
    p_pb.add_run('Strategy B — Tactical ATM Put Hedge (Stage 1 or 2 active)').bold = True
    p_pb.add_run(
        f'Trigger: 2-of-3 signals (VIX>20, Nifty<100-DMA by -2%, FPI net outflow). '
        f'Buy ATM Nifty Puts on {pb*100:.0f}% portfolio notional. '
        'Size to the v3.2 Beta-Adjusted Target %. Hold 1 month, roll if signals persist.'
    )
    p_pb.add_run('Strategy C — Crisis Short Futures (Stage 3 / VIX > 28)').bold = True
    p_pb.add_run(
        f'Short Nifty Futures covering {pb*100:.0f}% notional. '
        'Combine with collars to reduce cost. Never exceed 75% notional short.'
    )

    # 7. Dual-Path Comparison (Harvey Ch.3)
    _vix_w  = live_nifty.get('vix', 18.5)    if live_nifty else 18.5
    _dd_w   = live_nifty.get('gap_pct', 0.0) if live_nifty else 0.0
    eq_wt_w = metrics.get('equity_weight', 60.0)
    _pw     = compute_equity_reduction_path(pb, _vix_w, _dd_w, eq_wt_w, pv)

    doc.add_heading('VOLATILITY TARGETING — TWO-PATH COMPARISON  (Harvey Ch.3)', level=1)
    p_regime = doc.add_paragraph()
    p_regime.add_run(f"Current Regime: {_pw['regime']}  |  VIX: {_pw['vix']:.1f}  |  DD from Peak: {_pw['dd_pct']:.1f}%\n")

    doc.add_heading(f"Path 1 — {_pw['p1_label']}", level=2)
    p_p1 = doc.add_paragraph()
    p_p1.add_run(f"{_pw['p1_descriptor']}\n")
    p_p1.add_run(f"Equity Exposure  : {_pw['p1_equity_pct']:.1f}% (unchanged)\n")
    p_p1.add_run(f"Portfolio Beta   : {_pw['p1_beta']:.3f} (unchanged)\n")
    put_str = f"Put Allocation   : {_pw['p1_hedge_pct']:.0f}% of notional"
    if pv and 'p1_put_notional_inr' in _pw:
        put_str += f"  ({local_fmt_inr(_pw['p1_put_notional_inr'])})"
    p_p1.add_run(f"{put_str}\n")
    drag_str1 = f"Monthly Drag     : ~{_pw['p1_monthly_drag']:.2f}%"
    if pv and 'p1_monthly_drag_inr' in _pw:
        drag_str1 += f"  ({local_fmt_inr(_pw['p1_monthly_drag_inr'])})"
    p_p1.add_run(f"{drag_str1}\n")
    p_p1.add_run(f"Best For         : {_pw['p1_best_for']}\n")
    p_p1.add_run(f"Trade-off        : {_pw['p1_tradeoff']}\n")

    doc.add_heading(f"Path 2 — {_pw['p2_label']}", level=2)
    p_p2 = doc.add_paragraph()
    p_p2.add_run(f"{_pw['p2_descriptor']}\n")
    eq_str = f"Equity Trim      : -{_pw['p2_equity_trim']:.0f}% from {eq_wt_w:.1f}%  →  {_pw['p2_new_equity']:.1f}%"
    if pv and 'p2_equity_sell_inr' in _pw:
        eq_str += f"  (sell approx. {local_fmt_inr(_pw['p2_equity_sell_inr'])})"
    p_p2.add_run(f"{eq_str}\n")
    p_p2.add_run(f"Effective Beta   : {_pw['p2_effective_beta']:.3f} (reduced from {pb:.3f})\n")
    res_str = f"Residual Put     : {_pw['p2_hedge_pct']:.0f}% of notional"
    if pv and 'p2_put_notional_inr' in _pw:
        res_str += f"  ({local_fmt_inr(_pw['p2_put_notional_inr'])})"
    p_p2.add_run(f"{res_str}\n")
    p_p2.add_run(f"Put Drag/month   : ~{_pw['p2_monthly_drag']:.2f}%\n")
    p_p2.add_run(f"Opp. Cost/month  : ~{_pw['p2_opp_cost_monthly']:.2f}%  (forgone equity upside)\n")
    p_p2.add_run(f"Total Drag/month : ~{_pw['p2_total_drag']:.2f}%\n")
    p_p2.add_run(f"Best For         : {_pw['p2_best_for']}\n")
    p_p2.add_run(f"Trade-off        : {_pw['p2_tradeoff']}\n")

    p_note = doc.add_paragraph()
    p_note.add_run(
        "NOTE: Path 2 is adapted from Harvey's continuous vol-scalar (base × target_vol / current_vol) "
        "to a regime-based step function tied to VIX thresholds — avoiding daily rebalancing churn, "
        "STT, and capital gains friction for retail Indian investors."
    ).italic = True

    doc.save(output_path)



# ══════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description="Portfolio Hedge Engine v2.0")
    parser.add_argument("--file",  default="portfolio_input.csv",
                        help="Path to your portfolio CSV (default: portfolio_input.csv)")
    parser.add_argument("--value", type=float, default=None,
                        help="Total portfolio value in ₹ (optional). E.g. --value 5000000 for ₹50L")
    parser.add_argument("--output-dir", default="./portfolio_output",
                        help="Output directory (default: ./portfolio_output)")
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)
    portfolio_file = args.file

    print("\n" + "═"*60)
    print("  PORTFOLIO HEDGE ENGINE  v2.0 (Live Data Edition)")
    print("═"*60)

    if not os.path.exists(portfolio_file):
        print(f"\n  Portfolio file '{portfolio_file}' not found.")
        print("  Generating sample portfolio template...")
        generate_template_csv(portfolio_file)

    print(f"\n  Loading portfolio: {portfolio_file}")
    df_port = load_portfolio(portfolio_file)
    print(f"  Holdings: {len(df_port)} | Total allocation: {df_port['Allocation_Pct'].sum():.1f}%")

    metrics       = compute_portfolio_metrics(df_port, args.value)
    hedge_sizing  = compute_hedge_sizing(metrics["portfolio_beta"], args.value)
    scenario_df   = compute_scenario_analysis(metrics["portfolio_beta"], args.value)
    
    # --- NEW LIVE DATA FETCH ---
    live_nifty = fetch_live_nifty_status()
    # ---------------------------

    print(f"\n  Portfolio Beta: {metrics['portfolio_beta']:.3f}")
    if args.value:
        print(f"  Portfolio Value: {fmt_inr(args.value)}")

    xlsx_path = os.path.join(args.output_dir, "portfolio_hedge_analysis.xlsx")
    txt_path  = os.path.join(args.output_dir, "portfolio_hedge_report.txt")

    print("\n  Building Excel workbook (5 sheets)...")
    build_excel(df_port, metrics, hedge_sizing, scenario_df, xlsx_path)

    print("  Building text report...")
    generate_text_report(metrics, hedge_sizing, scenario_df, txt_path, live_nifty)

    print("═"*60 + "\n")

# ══════════════════════════════════════════════════════════════
# MAGNITUDE HEDGING v3.2 ENGINE
# ══════════════════════════════════════════════════════════════
import math

def sigmoid(x):
    return 1 / (1 + math.exp(-x))

def get_continuous_notional(score):
    # v3.2 FIX: Floor set to 10. Multipliers set to 30. (10 + 30 + 30 + 30 = 100 Max)
    n = 10.0
    n += 30.0 * sigmoid((score - 25.0) / 3.5)
    n += 30.0 * sigmoid((score - 50.0) / 3.5)
    n += 30.0 * sigmoid((score - 70.0) / 3.5)
    return min(100.0, max(10.0, n))

def norm_fpi(fpi):
    # v3.1 BUG 3 FIX: Denominator expanded to 8000 Cr. Inflows capped at -50.
    if fpi < 0:
        return min(100.0, (-fpi) / 8000.0 * 100.0) # Outflow
    else:
        return max(-50.0, (-fpi) / 8000.0 * 100.0) # Inflow

def compute_equity_reduction_path(portfolio_beta: float, vix: float, dd_pct: float,
                                   equity_weight_pct: float, portfolio_value: float = None) -> dict:
    """
    Computes the Harvey Ch.3 Vol-Targeting Path 2 recommendation:
    Regime-based equity reduction + residual hedge sizing.

    Rather than the continuous daily scalar (base × target_vol / current_vol),
    which causes weekly churn and tax friction for retail investors, this uses
    the same VIX thresholds already established in the v3.2 engine.

    Regime table (aligned with engine stage boundaries):
      VIX < 15   → Stage 1 low    : Hold equity, no change
      VIX 15-20  → Stage 1 elevated: Trim equity ~10%, lighter put
      VIX 20-28  → Stage 2        : Trim equity ~20%, smaller put
      VIX > 28   → Stage 3        : Full equity reduction + minimal put

    Returns a dict with both Path 1 (current) and Path 2 (equity reduction)
    side-by-side for comparison in dashboard and reports.
    """
    # ── Determine regime ──────────────────────────────────────────────────────
    if vix < 15:
        regime          = "Stage 1 — Low Volatility"
        equity_trim_pct = 0.0
        p2_hedge_pct    = 10.0   # minimal residual backstop
        regime_action   = "Hold equity, no change needed"
        urgency         = "low"
    elif vix < 20:
        regime          = "Stage 1 — Elevated Volatility"
        equity_trim_pct = 10.0
        p2_hedge_pct    = 15.0   # lighter put since equity already trimmed
        regime_action   = "Trim equity by ~10%, add lighter put as residual"
        urgency         = "watch"
    elif vix < 28:
        regime          = "Stage 2 — Active Drawdown"
        equity_trim_pct = 20.0
        p2_hedge_pct    = 20.0   # smaller put, equity reduction does the heavy lifting
        regime_action   = "Trim equity by ~20%, smaller put as backstop"
        urgency         = "active"
    else:
        regime          = "Stage 3 — Crisis"
        equity_trim_pct = 35.0
        p2_hedge_pct    = 10.0   # minimal put, equity reduction is primary action
        regime_action   = "Full equity reduction, minimal put as residual backstop"
        urgency         = "crisis"

    # ── Path 1 — current engine (Hedge in Place) ─────────────────────────────
    # Simple linear estimate of put allocation from VIX + dd_pct
    dd_abs = abs(dd_pct)
    p1_hedge_raw = 12.5 + 0.52 * vix + 0.48 * dd_abs + 5.10 * (vix * dd_abs / 100.0)
    p1_hedge_pct = round(min(100.0, max(10.0, p1_hedge_raw)), 1)
    p1_monthly_drag = round(p1_hedge_pct * 0.012, 2)   # ~1.2% per 1% put allocation / month
    p1_net_equity_exp = round(equity_weight_pct, 1)     # unchanged

    # ── Path 2 — equity reduction + residual put ─────────────────────────────
    new_equity_pct    = round(max(0.0, equity_weight_pct - equity_trim_pct), 1)
    effective_beta_p2 = round(portfolio_beta * (new_equity_pct / max(equity_weight_pct, 1.0)), 3)
    p2_monthly_drag   = round(p2_hedge_pct * 0.012, 2)

    # Combined drag: put cost + implicit opportunity cost of reduced equity exposure
    # Opportunity cost proxy: 1% per 10% equity trimmed (conservative bull-mkt assumption)
    p2_opportunity_cost = round(equity_trim_pct * 0.10, 2)  # per month, annualise ×12 for report
    p2_total_drag_monthly = round(p2_monthly_drag + p2_opportunity_cost, 2)

    result = {
        # Regime context
        "vix"            : round(vix, 1),
        "dd_pct"         : round(dd_pct, 2),
        "regime"         : regime,
        "urgency"        : urgency,

        # Path 1 — Hedge in Place
        "p1_label"       : "Hedge in Place",
        "p1_descriptor"  : "Keep your current allocation. Add derivative protection sized to your beta.",
        "p1_hedge_pct"   : p1_hedge_pct,
        "p1_equity_pct"  : p1_net_equity_exp,
        "p1_beta"        : round(portfolio_beta, 3),
        "p1_monthly_drag": p1_monthly_drag,
        "p1_action"      : f"Buy ATM puts at {p1_hedge_pct:.0f}% notional. Equity unchanged.",
        "p1_best_for"    : "Investors who cannot or prefer not to sell equity (tax, lock-in, conviction).",
        "p1_tradeoff"    : "Higher put premium. Full equity upside retained if hedge is wrong.",

        # Path 2 — Equity Reduction
        "p2_label"       : "Equity Reduction",
        "p2_descriptor"  : "Trim equity exposure first. Use a smaller hedge as a residual backstop.",
        "p2_equity_trim" : equity_trim_pct,
        "p2_new_equity"  : new_equity_pct,
        "p2_hedge_pct"   : p2_hedge_pct,
        "p2_effective_beta": effective_beta_p2,
        "p2_monthly_drag": p2_monthly_drag,
        "p2_opp_cost_monthly": p2_opportunity_cost,
        "p2_total_drag"  : p2_total_drag_monthly,
        "p2_action"      : regime_action,
        "p2_best_for"    : "Investors comfortable rebalancing and who want lower derivative exposure.",
        "p2_tradeoff"    : "Misses rally if VIX spike is a false alarm. Lower put cost.",
    }

    # ── Optional ₹ outputs ────────────────────────────────────────────────────
    if portfolio_value:
        result["p1_put_notional_inr"]   = round(portfolio_value * p1_hedge_pct / 100, 0)
        result["p1_monthly_drag_inr"]   = round(portfolio_value * p1_monthly_drag / 100, 0)
        result["p2_equity_sell_inr"]    = round(portfolio_value * equity_trim_pct / 100, 0)
        result["p2_put_notional_inr"]   = round(portfolio_value * p2_hedge_pct / 100, 0)
        result["p2_monthly_drag_inr"]   = round(portfolio_value * p2_monthly_drag / 100, 0)

    return result


def calculate_v3_magnitude_hedge(vix, rv20d, fpi_net, gap_pct, current_price, state, ret_5d,
                                  new_calendar_day=True):
    """
    Evaluates v3.2 logic. Requires previous state dictionary to manage 
    the De-escalation Gate and Drawdown penalties.
    """
    # --- NEW: 3-Minute EMA for VIX Smoothing ---
    alpha = 0.5
    prev_ema_vix = state.get('ema_vix', vix)
    vix = (vix * alpha) + (prev_ema_vix * (1 - alpha))  # Overwrite with smoothed value

    # 1. Update Peak, Trough & Drawdown
    peak = max(state.get('peak_price', current_price), current_price)
    # Note: Live engine uses absolute positive percentage for dd_pct
    dd_pct = abs((current_price - peak) / peak) * 100 if peak > 0 else 0
    trough = min(state.get('trough_price', current_price), current_price)
    
    if dd_pct == 0:
        trough = current_price 

    if dd_pct > 0:
        days_in_dd = state.get('days_in_dd', 0) + (1 if new_calendar_day else 0)
    else:
        days_in_dd = 0

    # 2. Stage 1: Onset Score Calculation
    rvN  = min(100.0, max(0.0, (rv20d - 5.0) / 25.0 * 100.0))
    gapN = min(100.0, max(0.0, (-gap_pct) / 20.0 * 100.0))
    fpiN = norm_fpi(fpi_net)
    vixN = min(100.0, max(0.0, (vix - 10.0) / 30.0 * 100.0))

    msf_active = (rv20d > 18.0 and ret_5d > -3.0)
    fpi_weight = 0.10 * (0.8 if msf_active else 1.0)

    onset_score = (0.48 * rvN) + (0.25 * gapN) + (fpi_weight * fpiN) + (0.17 * vixN)
    s1_target = get_continuous_notional(onset_score)

    # 3. Stage 2: Active Phase Calculation
    penalty = min(0.08 * days_in_dd, 15.0)
    s2_target = 12.5 + (0.52 * vix) + (0.48 * dd_pct) + (5.10 * (vix * dd_pct / 100.0)) - penalty
    s2_target = min(90.0, max(10.0, s2_target))

    # 4. Transition Logic
    final_target = max(s1_target, s2_target)
    if s1_target >= s2_target:
        active_stage = "Stage 1 — Onset"
    else:
        active_stage = "Stage 2 — Active Phase"

    # --- BUG 2 FIX: Calculate live RR so we can snapshot it ---
    recovery_ratio_live = 0.0
    if peak - trough > 0:
        recovery_ratio_live = (current_price - trough) / (peak - trough)

    # --- Capture SOD Snapshots for Gate Logic ---
    if new_calendar_day or 'sod_natural_target' not in state:
        sod_natural_target = final_target
        sod_buffer = 5.0 + max(0.0, ((vix - 15.0) / 10.0) * 3.0)
        sod_dd_pct = dd_pct
        sod_recovery_ratio = recovery_ratio_live  # BUG 2 FIX: Freeze RR at the open
    else:
        sod_natural_target = state.get('sod_natural_target', final_target)
        sod_buffer = state.get('sod_buffer', 5.0 + max(0.0, ((vix - 15.0) / 10.0) * 3.0))
        sod_dd_pct = state.get('sod_dd_pct', dd_pct)
        sod_recovery_ratio = state.get('sod_recovery_ratio', recovery_ratio_live)

    # 5. Stage 3: Escalation Overrides (Hard Glass-Smash)
    # --- BUG 1 FIX: Wrap overrides in the 3-day shock window guard ---
    if days_in_dd >= 3:
        if dd_pct >= 20:
            final_target = max(final_target, 75.0)
            active_stage = "Stage 3: Escalated (MAJOR minimum)"
        if dd_pct >= 15 and (rv20d > 28 or vix > 28):
            final_target = 100.0
            active_stage = "Stage 3: Escalated (SEVERE)"
        if dd_pct >= 25:
            final_target = 100.0
            active_stage = "Stage 3: Escalated (SEVERE unconditional)"

    # 6. Fix 7: De-escalation Gate (Memory Check)
    low_vol_days = state.get('low_vol_days', 0)
    if new_calendar_day:
        if rv20d < 22 and vix < 20:
            low_vol_days += 1
        else:
            low_vol_days = 0

    incoming_prev_hedge = state.get('prev_hedge', 0)

    if dd_pct == 0:
        incoming_prev_hedge = 0
        low_vol_days = 0
        days_in_dd = 0

    final_target = round(final_target, 1)
    gate_threshold_check = sod_natural_target + sod_buffer

    if incoming_prev_hedge > gate_threshold_check and "Stage 3" not in active_stage:
        # Bypass gate entirely if SOD drawdown is trivial (< 2%)
        if sod_dd_pct < 2.0:
            pass  
        else:
            # Escape Route 1: Price Momentum (uses frozen SOD value to prevent mid-day wobble)
            cleared_by_price = (sod_recovery_ratio >= 0.40)
            
            # Escape Route 2: Volatility Collapse 
            cleared_by_vol = (low_vol_days >= 5)
            
            # Gate clears if EITHER condition is met
            gate_cleared = (cleared_by_price or cleared_by_vol) and (days_in_dd >= 3)
            
            if not gate_cleared:
                final_target = incoming_prev_hedge
                # Do NOT override active_stage — market conditions determine stage
                # Gate lock is surfaced separately via diagnostics

    new_prev_hedge = max(incoming_prev_hedge, final_target) if dd_pct > 0 else 0.0

    # 7. Compile Final State and Diagnostics
    new_state = {
        'peak_price': peak,
        'trough_price': trough,
        'days_in_dd': days_in_dd,
        'low_vol_days': low_vol_days,
        'prev_hedge': new_prev_hedge,
        'ema_vix': vix,
        'sod_dd_pct': sod_dd_pct,
        'sod_natural_target': sod_natural_target,
        'sod_buffer': sod_buffer,
        'sod_recovery_ratio': sod_recovery_ratio  # Added to state
    }

    # Compute gate_locked flag for diagnostics
    _gate_active = (
        incoming_prev_hedge > gate_threshold_check
        and sod_dd_pct >= 2.0
        and "Stage 3" not in active_stage
    )
    _gate_locked = _gate_active and not (
        (sod_recovery_ratio >= 0.40 or low_vol_days >= 5) and days_in_dd >= 3
    )

    diagnostics = {
        'onset_score'     : round(onset_score, 1),
        's1_target'       : round(s1_target, 1),
        's2_penalty'      : round(penalty, 1),
        's2_target'       : round(s2_target, 1),
        'current_dd'      : round(dd_pct, 2),
        'gate_locked'     : _gate_locked,
        'gate_prev_hedge' : round(incoming_prev_hedge, 1),
        'recovery_ratio'  : round(sod_recovery_ratio, 3),
        'low_vol_days'    : low_vol_days,
    }

    return final_target, active_stage, diagnostics, new_state

if __name__ == "__main__":
    main()
